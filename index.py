import os
import jwt
from flask import Flask, request, jsonify, send_file, session
import instaloader
from docx import Document
import datetime
import json
# from flask_pymongo import PyMongo
from flask_cors import CORS

app = Flask(__name__)
CORS(app)
app.secret_key = '1234'
JWT_SECRET = 'karthik1234'  # Replace with a strong, random key
JWT_ALGORITHM = 'HS256'
# MongoDB configuration
# app.config["MONGO_URI"] = "mongodb+srv://amrindersingh292004:YrUO6w88OMH3piB3@clusterinstascrapper.stjux.mongodb.net/"
# mongo = PyMongo(app)

@app.route('/login', methods=['POST'])
def login():
    data = request.get_json()

    # Extract the username and password from the request data
    username = data.get('username')
    password = data.get('password')

    # Validate username and password
    if username == 'admin' and password == 'Parse1234':
        # Generate a JWT token that expires in 1 hour
        token = jwt.encode({
            'username': username,
            'exp': datetime.datetime.now() + datetime.timedelta(hours=1)
        }, JWT_SECRET, algorithm=JWT_ALGORITHM)

        return jsonify({'message': 'Login successful', 'token': token}), 200
    else:
        return jsonify({'message': 'Invalid username or password'}), 401


@app.route('/protected', methods=['GET'])
def protected():
    auth_header = request.headers.get('Authorization')

    if not auth_header:
        return jsonify({'message': 'Token is missing'}), 401

    try:
        # Extract the token from the Authorization header and decode it
        decoded_token = jwt.decode(auth_header, JWT_SECRET, algorithms=[JWT_ALGORITHM])

        # If the token is valid, return the protected resource
        return jsonify({'message': f'Welcome, {decoded_token["username"]}!'}), 200

    except jwt.ExpiredSignatureError:
        return jsonify({'message': 'Token has expired'}), 401
    except jwt.InvalidTokenError:
        return jsonify({'message': 'Invalid token'}), 401


@app.route('/logout', methods=['POST'])
def logout():
    return jsonify({'message': 'Logout successful'}), 200

@app.route('/instagram', methods=['POST'])
def insta_scraper():

    auth_header = request.headers.get('Authorization')
    if not auth_header:
        return jsonify({"msg" : "u r not logged in"}) , 401
    try:
        decoded_token = jwt.decode(auth_header, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        return jsonify({"message": "Token has expired"}), 401
    except jwt.InvalidTokenError:
        return jsonify({"message": "Invalid token"}), 401
    
    if decoded_token:
        name = request.json.get('name')
        pwd = request.json.get('pwd')

        if not name or not pwd:
            return jsonify({'error': 'Username and password are required'}), 400

        try:
            # Initialize Instaloader
            L = instaloader.Instaloader()
            print("before insta login")
            # Login to Instagram
            L.login(name, pwd)
            print("after login login")
            # Define the profile to scrape
            profile_name = name
            profile = instaloader.Profile.from_username(L.context, profile_name)

            # Create a new Document
            doc = Document()
            doc.add_heading(f'Instagram Profile Data: {profile_name}', 0)

            # Add profile information to the document
            doc.add_heading('Profile Information', level=1)
            doc.add_paragraph(f'Username: {profile.username}')
            doc.add_paragraph(f'Full Name: {profile.full_name}')
            doc.add_paragraph(f'Biography: {profile.biography}')
            doc.add_paragraph(f'Followers: {profile.followers}')
            doc.add_paragraph(f'Following: {profile.followees}')
            doc.add_paragraph(f'Number of Posts: {profile.mediacount}')
            # fetching saved post
            doc.add_heading('Saved Posts', level=1)
            saved_posts = list(profile.get_saved_posts())

            for post in saved_posts:
                doc.add_paragraph(f"Caption: {post.caption}")
                doc.add_paragraph(f'Likes: {post.likes}')
                doc.add_paragraph(f'Comments: {post.comments}')
                if post.video_url:
                    doc.add_paragraph(f'Video URL: {post.video_url}')
                else:
                    doc.add_paragraph(f'Image URL: {post.url}')
                doc.add_paragraph('')
            # fetching similar account 
            doc.add_heading(f'similar accounts', level=1)
            similar_account  = profile.get_similar_accounts()
            for account in similar_account:
                doc.add_paragraph(f'username : {account.username}')
                doc.add_paragraph(f'name : {account.full_name}')
                doc.add_paragraph(f'followers : {account.followers}')
                doc.add_paragraph(f'folloing : {account.followees}')


            # Add recent posts to the document
            doc.add_heading('Recent Posts', level=1)
            posts = list(profile.get_posts())
            for post in posts:
                doc.add_heading(post.date.strftime('%Y-%m-%d'), level=2)
                doc.add_paragraph(f"Caption: {post.caption}")
                doc.add_paragraph(f'Likes: {post.likes}')
                doc.add_heading(f'Comments: ', level=2)
                comments = post.get_comments()
                for comment in comments:
                    doc.add_paragraph(f'{comment.owner.username} : {comment.text}')
                if post.video_url:
                    doc.add_paragraph(f'Video URL: {post.video_url}')
                else:
                    doc.add_paragraph(f'Image URL: {post.url}')
                doc.add_paragraph('')

            # Save the document
            doc.save(f'./ScrappedFiles/{profile.username}.docx')

            # Save the data to MongoDB
            # instagram_data = {
            #     'username': profile.username,
            #     'full_name': profile.full_name,
            #     'biography': profile.biography,
            #     'followers': profile.followers,
            #     'following': profile.followees,
            #     'mediacount': profile.mediacount,
            #     'posts': []
            # }
            # for post in profile.get_posts():
            #     instagram_data['posts'].append({
            #         'date': post.date,
            #         'url': post.url,
            #         'caption': post.caption,
            #         'likes': post.likes,
            #         'comments': post.comments,
            #         'video_url': post.video_url
            #     })
            # mongo.db.instagram_profiles.insert_one(instagram_data)
            return jsonify({'message': 'Data scraped and saved to instagram_profile_data.docx and MongoDB'})

        except instaloader.exceptions.BadCredentialsException:
            return jsonify({'error': 'Invalid username or password'}), 401
        except Exception as e:
            return jsonify({'error': f'An error occurred: {e}'}), 500
    else:
        return jsonify({"smtg went wrong"}) , 401

@app.route('/instagram/download', methods=['GET'])
def download_docx():
    # Check if the user is logged in
    # if 'username' not in session:
    #     return jsonify({'error': 'You are not logged in'}), 401
    username = request.args.get('username')
    # Check if the DOCX file exists
    file_path = f'./ScrappedFiles/{username}.docx'
    print("after path")
    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404

    # Send the file as a download response
    print("sending atach")
    return send_file(file_path, as_attachment=True, download_name=f'{username}.docx')


if __name__=="__main__":
    app.run(debug=True,port=5000)